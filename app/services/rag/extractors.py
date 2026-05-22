"""Text extraction helpers for MVP RAG ingestion."""

import json
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

RowDict = dict[str, object]

SUPPORTED_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".pdf", ".docx", ".xlsx"}
SUPPORTED_CONTENT_TYPES = {
    "text/plain",
    "text/markdown",
    "text/csv",
    "application/json",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    # Windows Explorer and some browsers report this type for .xlsx files.
    "application/vnd.ms-excel",
}

MAX_XLSX_SHEETS = 5
MAX_XLSX_ROWS_PER_SHEET = 300


def extract_text(path: str | Path, filename: str, content_type: str) -> RowDict:
    """Extract text and metadata from a supported text-like file."""
    file_path = Path(path)
    extension = Path(filename).suffix.lower()
    normalized_content_type = content_type.strip().lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError("unsupported file type")
    if normalized_content_type and normalized_content_type not in SUPPORTED_CONTENT_TYPES:
        raise ValueError("unsupported file type")

    if extension == ".pdf":
        text = _extract_pdf_text(file_path)
    elif extension == ".docx":
        text = _extract_docx_text(file_path)
    elif extension == ".xlsx":
        text = _extract_xlsx_text(file_path)
    else:
        raw_text = _read_text_with_fallback(file_path)
        if extension == ".json":
            text = _normalize_json_text(raw_text)
        else:
            text = raw_text

    return {
        "text": text,
        "metadata": {
            "filename": filename,
            "content_type": normalized_content_type,
            "extension": extension,
            "size_bytes": file_path.stat().st_size,
            "line_count": _count_lines(text),
        },
    }


def _read_text_with_fallback(path: Path) -> str:
    data = path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _normalize_json_text(text: str) -> str:
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        return text
    return json.dumps(parsed, ensure_ascii=False, indent=2, sort_keys=True)


def _extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        page_text = []
        for page in reader.pages:
            page_text.append(page.extract_text() or "")
        return "\n".join(text for text in page_text if text).strip()
    except Exception as exc:
        raise ValueError("file text extraction failed") from exc


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document

        document = Document(str(path))
        blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells if cell.text]
                if cells:
                    blocks.append("\t".join(cells))
        return "\n".join(blocks).strip()
    except Exception as exc:
        raise ValueError("file text extraction failed") from exc


def _extract_xlsx_text(path: Path) -> str:
    import io
    import zipfile

    try:
        import openpyxl

        # Pass BytesIO so openpyxl skips its file-extension check.
        # Storage paths are extensionless ("source"), which openpyxl rejects
        # via InvalidFileException when given a string path.
        workbook = openpyxl.load_workbook(io.BytesIO(path.read_bytes()), read_only=True, data_only=True)
        blocks: list[str] = []
        for sheet_name in workbook.sheetnames[:MAX_XLSX_SHEETS]:
            sheet = workbook[sheet_name]
            rows_text: list[str] = []
            row_count = 0
            for row in sheet.iter_rows(values_only=True):
                if row_count >= MAX_XLSX_ROWS_PER_SHEET:
                    break
                cells = [str(cell).strip() if cell is not None else "" for cell in row]
                if not any(cells):
                    continue
                rows_text.append(" | ".join(cells))
                row_count += 1
            if rows_text:
                blocks.append(f"[Sheet: {sheet_name}]")
                blocks.extend(rows_text)
        workbook.close()
        return "\n".join(blocks).strip()
    except zipfile.BadZipFile as exc:
        raise ValueError(
            "File is not a valid .xlsx file. "
            "If this is an older .xls file, open it in Excel and save as .xlsx."
        ) from exc
    except Exception as exc:
        logger.warning(
            "xlsx_extraction_failed",
            extra={"path": str(path), "error_type": type(exc).__name__, "error": str(exc)},
        )
        raise ValueError("file text extraction failed") from exc


def _count_lines(text: str) -> int:
    if not text:
        return 0
    return text.count("\n") + 1
